
import os
from pathlib import Path
from typing import Dict, Optional

from docx import Document as DocxDocument
import pdfminer.high_level

import asyncio
import httpx
import json
import re

import matplotlib.pyplot as plt
import matplotlib.patches as patches

META_PATTERNS = [
    r"^Použij.*formát", r"^Nepopisuj", r"^Vytvoř tabulku", r"^Uveď",
    r"^Text napiš", r"^Zdrojový text", r"^Vrať jen", r"^Vypiš"
]

def _strip_meta(text: str) -> str:
    lines = text.splitlines()
    cleaned = [l for l in lines if not any(re.search(p, l.strip(), flags=re.I) for p in META_PATTERNS)]
    # když by náhodou zmizely všechny řádky, vrať původní
    return "\n".join(cleaned) if any(s.strip() for s in cleaned) else text

USE_AZURE_SPEECH = bool(os.getenv("AZURE_SPEECH_KEY"))

class GenerationError(Exception):
    pass

def _read_text_from_path(p: Path) -> str:
    try:
        if p.suffix.lower() == ".docx":
            d = DocxDocument(str(p))
            return "\n".join([para.text for para in d.paragraphs])
        if p.suffix.lower() == ".pdf":
            return pdfminer.high_level.extract_text(str(p))
        if p.suffix.lower() in [".txt", ".md"]:
            return p.read_text(encoding="utf-8", errors="ignore")
        return ""
    except Exception:
        return ""

def _collect_corpus(input_dir: Path) -> str:
    texts = []
    for p in input_dir.rglob("*"):
        if p.is_file():
            texts.append(_read_text_from_path(p))
    corpus = "\n\n".join(t for t in texts if t.strip())
    if not corpus.strip():
        raise GenerationError("Nebyl nalezen žádný čitelný text (PDF/DOCX/TXT).")
    return corpus[:200000]

def _build_prompt(corpus: str, audience: str, style: str, attachments: list[str]) -> str:
    return f"""
Napiš hotový a konkrétní návod v češtině pro administrativní pracovníky veřejné správy.
Nepopisuj, co budeš dělat – rovnou napiš obsah. Neopakuj zadání ani slova jako
„použij, napiš, uveď, vytvoř“. Nevkládej žádné vysvětlování úkolu.

Vrať jen tyto sekce s vyplněným obsahem (bez instrukcí):

# Metodický převodník
(zjednodušený výklad pro administrativní pracovníky)

## Úvod
2–5 vět o účelu a působnosti metodiky (bez metatextu).

## Přehled fází
V bodech uveď skutečné fáze/kroky procesu z textu. Ke každé jedna věta, co se děje.

## Kroky a odpovědnosti
Tabulka v markdownu:
| Kdo | Co | Jak | Do kdy | Přílohy |
Vyplň reálnými informacemi. „Přílohy“ vybírej jen z těchto názvů (pokud dávají smysl pro daný krok):
{attachments}

## Na co si dát pozor
5–8 konkrétních rizik/kontrolních bodů z metodiky.

Zdrojový materiál (použij pro fakta, pojmy, role, lhůty a přílohy):
{corpus}
"""

async def _call_llm(prompt: str) -> str:
    openai_key = os.getenv("OPENAI_API_KEY")
    azure_key = os.getenv("AZURE_OPENAI_KEY")
    if azure_key and os.getenv("AZURE_OPENAI_ENDPOINT") and os.getenv("AZURE_OPENAI_DEPLOYMENT"):
        endpoint = os.getenv("AZURE_OPENAI_ENDPOINT").rstrip("/")
        deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT")
        url = f"{endpoint}/openai/deployments/{deployment}/chat/completions?api-version=2024-02-15-preview"
        headers = {"api-key": azure_key, "Content-Type": "application/json"}
        payload = {
            "messages": [
                {"role": "system", "content": (
                  "Jsi český úřední asistent. Tvoje odpověď NESMÍ obsahovat popisy typu "
                  "‘vytvoř, vypiš, použij, v tomto dokumentu’. Rovnou napiš hotový obsah: "
                  "zejména konkrétní fáze, kroky, role, termíny a přílohy. "
                  "Piš stručně, úředně a srozumitelně, bez anglicismů."
                )},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.1,
            "top_p": 0.8,
            "max_tokens": 3200,
        }
        async with httpx.AsyncClient(timeout=60) as client:
            r = await client.post(url, headers=headers, json=payload)
            r.raise_for_status()
            data = r.json()
            return data["choices"][0]["message"]["content"]
    elif openai_key:
        headers = {"Authorization": f"Bearer {openai_key}", "Content-Type": "application/json"}
        payload = {
            "model": "gpt-4o-mini",
            "messages": [
                {"role": "system", "content": (
                  "Jsi český úřední asistent. Tvoje odpověď NESMÍ obsahovat popisy typu "
                  "‘vytvoř, vypiš, použij, v tomto dokumentu’. Rovnou napiš hotový obsah: "
                  "zejména konkrétní fáze, kroky, role, termíny a přílohy. "
                  "Piš stručně, úředně a srozumitelně, bez anglicismů."
                )},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.1,
            "top_p": 0.8,
            "max_tokens": 3200,
        }
        async with httpx.AsyncClient(timeout=60) as client:
            r = await client.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
            r.raise_for_status()
            data = r.json()
            return data["choices"][0]["message"]["content"]
    else:
        return "POZNÁMKA: LLM není nakonfigurováno. Uveďte API klíče. Náhled shrnutí:\n\n" + prompt[:1500]

def _make_docx(text: str, out_path: Path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    r = title.add_run("Metodický pokyn\n(zjednodušený výklad pro personální útvary)")
    r.bold = True
    r.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in text.splitlines():
        if line.strip().startswith("# "):
            doc.add_heading(line.strip("# ").strip(), level=1)
        elif line.strip().startswith("## "):
            doc.add_heading(line.strip("# ").strip(), level=2)
        else:
            doc.add_paragraph(line)
    doc.save(out_path)

def _make_png_from_phases(output_dir: Path, phases: list[str]) -> str:
    png_path = output_dir / "schema.png"
    import matplotlib.pyplot as plt
    import matplotlib.patches as patches

    steps = phases[:14] if phases else ["Fáze 1", "Fáze 2"]
    fig, ax = plt.subplots(figsize=(9, max(6, 1.1*len(steps))))
    ax.axis('off')
    W, H = 0.82, 0.07; x = 0.09
    ys = [0.92 - i*0.08 for i in range(len(steps))]
    for i, (y, text) in enumerate(zip(ys, steps)):
        rect = patches.FancyBboxPatch((x, y), W, H,
                                      boxstyle="round,pad=0.02,rounding_size=0.02",
                                      linewidth=1, edgecolor="black", facecolor="white")
        ax.add_patch(rect)
        ax.text(x + W/2, y + H/2, text, ha='center', va='center',
                wrap=True, fontsize=10)
        if i < len(ys)-1:
            ax.annotate("", xy=(x+W/2, y),
                        xytext=(x+W/2, ys[i+1]+H),
                        arrowprops=dict(arrowstyle="->", lw=1))
    fig.savefig(png_path, dpi=200, bbox_inches="tight")
    return str(png_path)

def _make_tts(text: str, out_path: Path) -> Optional[str]:
    if not USE_AZURE_SPEECH:
        return None
    try:
        import azure.cognitiveservices.speech as speechsdk
        speech_config = speechsdk.SpeechConfig(subscription=os.getenv("AZURE_SPEECH_KEY"), region=os.getenv("AZURE_SPEECH_REGION", "westeurope"))
        speech_config.speech_synthesis_language = "cs-CZ"
        speech_config.speech_synthesis_voice_name = os.getenv("AZURE_SPEECH_VOICE", "cs-CZ-AntoninNeural")
        audio_config = speechsdk.audio.AudioOutputConfig(filename=str(out_path))
        synthesizer = speechsdk.SpeechSynthesizer(speech_config=speech_config, audio_config=audio_config)
        synthesizer.speak_text_async(text[:5000]).get()
        return str(out_path)
    except Exception:
        return None
        
def _list_attachments(input_dir: Path) -> list[str]:
     """Vrať jen názvy souborů (bez cest), aby je LLM mohlo mapovat k přílohám."""
     out = []
     for p in input_dir.rglob("*"):
         if p.is_file():
             out.append(p.name)
     return sorted(out)[:300]  # bezpečnostní limit

def _extract_phases_from_json_block(text: str) -> list[str]:
     """Hledá blok ```json ...``` s klíčem 'phases'. Vrátí seznam názvů fází."""
     m = re.search(r"```json\s*(\{.*?\})\s*```", text, flags=re.S)
     if not m:
         return []
     try:
         data = json.loads(m.group(1))
         phases = data.get("phases", [])
         if phases and isinstance(phases[0], dict) and "name" in phases[0]:
             phases = [x["name"] for x in phases]
         return [p.strip() for p in phases if isinstance(p, str) and p.strip()]
     except Exception:
         return []

def _first_bullets_as_phases(markdown_text: str, max_items: int = 12) -> list[str]:
     """Záložní metoda: vezmi první seznam odstavců pod '## Přehled fází'."""
     block = re.split(r"^##\s+Přehled fází.*?$", markdown_text, maxsplit=1, flags=re.I|re.M)
     if len(block) < 2:
         return []
     after = block[1]
     items = re.findall(r"^\s*[-*]\s*(.+)$", after, flags=re.M)
     # odstraň případný popis za pomlčkou
     return [re.sub(r"\s*[–-]\s*.*$", "", it).strip() for it in items][:max_items]

async def process_inputs_and_generate(input_dir: Path, output_dir: Path, audience: str, style: str) -> Dict[str, str]:
    corpus = _collect_corpus(input_dir)
    attachments = _list_attachments(input_dir)
    prompt = _build_prompt(corpus, audience, style, attachments)
    text = await _call_llm(prompt)
    text = _strip_meta(text)

    # Získat fáze pro schéma
    phases = _extract_phases_from_json_block(text)
    if not phases:
        phases = _first_bullets_as_phases(text)
    if not phases:
        phases = ["Příprava", "Zpracování", "Schválení", "Realizace", "Uzavření"]  # nouzová výplň


    docx_path = output_dir / "vystup_navod.docx"
    _make_docx(text, docx_path)
    png_path = _make_png_from_phases(output_dir, phases)
    audio_path = _make_tts(text, output_dir / "shrnutí.mp3")

    return {"docx": str(docx_path), "png": str(png_path), "audio": audio_path}
